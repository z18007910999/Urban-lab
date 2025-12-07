# app.py — Urban Lab · News Categorizer (manual review; only write to news_reviews)
# deps: streamlit supabase==2.* python-dotenv pandas

import pandas as pd
import streamlit as st
from datetime import date, timedelta, datetime, timezone
from supabase_io import fetch_articles, supabase  # 复用你的封装与客户端
from io import BytesIO
import os, requests, re  # ✅ changed: add re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from html import escape
from PIL import Image
import streamlit.components.v1 as components


# ===== Weekly DOCX helpers =====

OUTPUT_DIR = r"D:\Python Project\weekly outcome"  # 目标保存目录（可在UI里改）

def start_of_week(d: date) -> date:
    return d - timedelta(days=d.weekday())  # 周一

def end_of_week(start: date) -> date:
    return start + timedelta(days=6)

def _add_label_value(doc: Document, label: str, value: str, bold_label=True):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} ")
    r1.bold = bold_label
    r1.font.size = Pt(11)
    r2 = p.add_run(value or "")
    r2.font.size = Pt(11)
    return p

def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url,
                          reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# =========================
# Image fetching utilities
# =========================

# 统一的远程图片下载（requests + PIL 验证）
def fetch_remote_img(url: str) -> bytes | None:
    """
    下载远程图片并返回二进制；若不是图片或失败则返回 None。
    """
    if not url:
        return None
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        # 用 PIL 验证是否为图片
        img = Image.open(BytesIO(resp.content))
        img.verify()
        return resp.content
    except Exception:
        return None

# ✅ changed: 你提供的 curl 头和 cookie（可按需更新）
NYT_HEADERS = {
    
def fetch_og_image_url_with_curl(page_url: str) -> str | None:
    """
    用 curl 等价的 headers + cookie 抓页面，解析 og:image。
    """
    if not page_url:
        return None
    try:
        headers = NYT_HEADERS.copy()
        headers["Cookie"] = NYT_COOKIE
        resp = requests.get(page_url, headers=headers, timeout=12, allow_redirects=True)
        resp.raise_for_status()
        html = resp.text
        m = re.search(
            r'<meta[^>]+property=["\']og:image["\'][^>]*content=["\']([^"\']+)["\']',
            html, flags=re.I
        )
        if not m:
            return None
        og = m.group(1).strip()
        if og.startswith("//"):
            og = "https:" + og
        return og
    except Exception:
        return None

@st.cache_data(show_spinner=False, ttl=3600)
def _fetch_reviews_week(monday: date):
    """从审核表读取本周 [Mon..Sun] 的记录；表名优先 news_reviews，回退 '\"News_reviews\"'。"""
    start_s, end_s = monday.isoformat(), end_of_week(monday).isoformat()
    table_candidates = ["news_reviews", '"News_reviews"']
    last_err = None
    for tbl in table_candidates:
        try:
            res = (
                supabase.table(tbl)
                .select("*")
                .gte("publish_date", start_s)
                .lte("publish_date", end_s)
                .order("publish_date", desc=False)
                .execute()
            )
            data = res.data or []
            if data is not None:
                return data
        except Exception as e:
            last_err = e
            continue
    raise RuntimeError(f"Read reviews failed: {last_err}")

def build_weekly_docx(rows: list[dict], monday: date, author: str) -> BytesIO:
    """按模板生成 DOCX 并返回字节缓冲（供下载/另存）。"""
    week_text = monday.strftime("%B %d, %Y")  # e.g., October 27, 2025
    doc = Document()

    for i, r in enumerate(rows):
        if i > 0:
            doc.add_page_break()

        # Week of
        p_week = doc.add_paragraph()
        run = p_week.add_run(f"Week of {week_text}")
        run.bold = True; run.font.size = Pt(12)

        # Title
        _add_label_value(doc, "Title:", r.get("title",""))

        # Source / Date Published / Link / Author
        _add_label_value(doc, "Source:", r.get("publisher",""))
        pubdate = r.get("publish_date")
        pubdate_str = ""
        try:
            if pubdate:
                pubdate_str = pd.to_datetime(pubdate).strftime("%m.%d.%Y")
        except Exception:
            pass
        _add_label_value(doc, "Date Published:", pubdate_str)

        p_link = doc.add_paragraph()
        r_label = p_link.add_run("Link: "); r_label.bold = True; r_label.font.size = Pt(11)
        link = r.get("link") or r.get("url") or ""
        if link:
            _add_hyperlink(p_link, link, link)

        _add_label_value(doc, "Urban Lab Author:", author)

        # ✅ changed: Article Photograph，优先 image_url，再尝试用 curl 头从文章页抓 og:image
        image_url = (r.get("image_url") or "").strip()
        img_bytes = None
        if image_url:
            img_bytes = fetch_remote_img(image_url)
        if (not img_bytes) and link:
            og_url = fetch_og_image_url_with_curl(link)
            if og_url:
                img_bytes = fetch_remote_img(og_url)

        if img_bytes:
            doc.add_paragraph("Article Photograph:")
            try:
                doc.add_picture(BytesIO(img_bytes), width=Inches(6.5))
            except Exception:
                doc.add_paragraph("")
        else:
            _add_label_value(doc, "Article Photograph:", "")

        # Summary
        p_sum = doc.add_paragraph()
        r1 = p_sum.add_run("Article Summary: "); r1.bold = True; r1.font.size = Pt(11)
        p_sum.add_run(r.get("summary","")).font.size = Pt(11)

        # Initiatives（红色/加粗）
        p_init = doc.add_paragraph()
        r2 = p_init.add_run("Initiative: "); r2.bold = True; r2.font.size = Pt(11)
        r3 = p_init.add_run(r.get("categories","")); r3.bold = True; r3.font.size = Pt(11)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def build_weekly_block_text(row: dict, categories: str, author: str) -> str:
    """
    生成一个纯文本版的 Weekly Report 块，用于在页面上直接复制到 Google Docs。
    """
    pubdate = row.get("publish_date")
    # 计算 Week of
    monday_text = ""
    try:
        if isinstance(pubdate, date):
            d = pubdate
        else:
            d = pd.to_datetime(pubdate).date()
        monday_text = start_of_week(d).strftime("%B %d, %Y")
    except Exception:
        pass

    # 格式尽量和 DOCX 保持一致
    lines = []
    if monday_text:
        lines.append(f"Week of {monday_text}")
        lines.append("")

    lines.append(f"Title: {row.get('title', '')}")
    lines.append(f"Source: {row.get('publisher', '')}")

    # Date Published 格式  MM.DD.YYYY
    pub_str = ""
    try:
        if isinstance(pubdate, date):
            pub_str = pubdate.strftime("%m.%d.%Y")
        elif pubdate:
            pub_str = pd.to_datetime(pubdate).strftime("%m.%d.%Y")
    except Exception:
        pub_str = str(pubdate or "")
    lines.append(f"Date Published: {pub_str}")

    link = row.get("url") or row.get("link") or ""
    lines.append(f"Link: {link}")

    lines.append(f"Urban Lab Author: {author}")
    lines.append("")
    lines.append("Article Summary:")
    lines.append(row.get("summary", "") or "")
    lines.append("")
    lines.append(f"Initiative: {categories or ''}")

    return "\n".join(lines)


st.set_page_config(page_title="Urban Lab · News Categorizer", page_icon="📰", layout="wide")
st.markdown("""
<style>
/* 保证所有列与 Markdown 容器不超出父容器 */
.block-container, .stMarkdown, .stColumn {
  max-width: 100% !important;
}

/* 让 Markdown 文本可以在任意位置换行，避免挤出右侧列 */
.stMarkdown p, .stMarkdown div, .stMarkdown span {
  overflow-wrap: anywhere !important;
  word-break: break-word !important;
  white-space: normal !important;
}

/* 列容器不允许把子元素“挤出” */
[data-testid="stVerticalBlock"] {
  overflow: hidden !important;
}

/* DataFrame 本身允许横向滚动，但不突破列宽 */
[data-testid="stDataFrame"] {
  max-width: 100% !important;
  overflow-x: auto !important;
}
</style>
""", unsafe_allow_html=True)
st.title("📰 Urban Lab — News Article Categorization")

CATEGORIES = [
    "Housing Affordability",
    "Culture Led Development",
    "Net Zero Cities",
    "Public/Private Development",
]

# ---------------------------
# 数据加载（从 News_storage）
# ---------------------------
@st.cache_data(show_spinner=False, ttl=300)
def load_articles(limit: int = 1000) -> pd.DataFrame:
    rows = fetch_articles(limit=limit)
    recs = []
    for r in rows:
        recs.append({
            "id":           r.get("id"),
            "title":        r.get("title", ""),
            "publisher":    r.get("Publisher") or r.get("creator",""),
            "publish_date": r.get("pubdate",""),
            "url":          r.get("link",""),
            "summary":      r.get("summary",""),
            "category":     r.get("Category") or "",
            "image_url":    r.get("image_url",""),
        })
    df = pd.DataFrame(recs)
    if not df.empty:
        df["publish_date"] = pd.to_datetime(df["publish_date"], errors="coerce").dt.date
    return df

df_all = load_articles()
if df_all.empty:
    st.info("Supabase 表 `News_storage` 暂无数据或读取失败。请检查环境变量和 RLS。")
    st.stop()

# ---------------------------
# 侧边栏筛选
# ---------------------------
with st.sidebar:
    st.subheader("Filters")

    min_d = df_all["publish_date"].min()
    max_d = df_all["publish_date"].max()
    if pd.isna(min_d) or pd.isna(max_d):
        from_d, to_d = None, None
    else:
        from_d = st.date_input("Start date", value=min_d, min_value=min_d, max_value=max_d)
        to_d   = st.date_input("End date",   value=max_d, min_value=min_d, max_value=max_d)

    all_pubs = sorted([p for p in df_all["publisher"].dropna().unique() if str(p).strip()])
    sel_pubs = st.multiselect("News Publisher", all_pubs, default=all_pubs)

    q = st.text_input("Search title/summary", value="", placeholder="type keywords").strip()
    only_unreviewed = st.toggle("Show only unreviewed (Category is NULL/empty)", value=False)

# 应用筛选
df = df_all.copy()
if from_d and to_d:
    df = df[(df["publish_date"].notna()) & (df["publish_date"] >= from_d) & (df["publish_date"] <= to_d)]
if sel_pubs:
    df = df[df["publisher"].isin(sel_pubs)]
if q:
    ql = q.lower()
    df = df[df.apply(lambda r: ql in (r["title"] or "").lower()
                              or ql in (r["summary"] or "").lower(), axis=1)]
if only_unreviewed:
    df = df[(df["category"].isna()) | (df["category"] == "")]

# ---------------------------
# 三列布局
# ---------------------------
left, mid, right = st.columns([3.2, 6, 3.2], gap="large")

# 左列：文章选择（只显示标题）
with left:
    st.subheader("Select Article")

    # ✅ 不再显示表格，只构建下拉选项（显示标题）
    options = [(int(r["id"]), r["title"]) for _, r in df[["id", "title"]].iterrows()]
    if not options:
        st.warning("当前筛选结果为空，请调整 Filters。")
        st.stop()

    current = st.selectbox(
        "Select an article",
        options,
        format_func=lambda x: x[1],   # 只显示标题
        index=0,
    )
    current_id = current[0]

# 当前文章（保持原写法）
row = df.set_index("id").loc[current_id].to_dict()

# ---------------------------
# 中列：审核面板（AI Pre-selection = Category）
# ---------------------------
with mid:
    st.subheader("News Categorizer")

    if row.get("url"):
        st.markdown(f"### {row['title']}  ↗")
    else:
        st.markdown(f"### {row['title']}")

    # ✅ changed: Article image 渲染逻辑
    shown_image = False

    # 1) 优先使用存储字段 image_url（直接当图片 URL 下载）
    img_url = (row.get("image_url") or "").strip()
    if img_url:
        img_bytes = fetch_remote_img(img_url)
        if img_bytes:
            st.image(img_bytes, caption="Article image", use_container_width=True)
            shown_image = True

    # 2) 回退：从文章页抓 og:image（使用你的 curl 头和 cookie）
    if not shown_image and row.get("url"):
        og_url = fetch_og_image_url_with_curl(row["url"])
        if og_url:
            img_bytes2 = fetch_remote_img(og_url)
            if img_bytes2:
                st.image(img_bytes2, caption="Article image", use_container_width=True)
                shown_image = True

    if not shown_image:
        st.info("No image available.")

    # --- Summary ---
    raw_summary = row.get("summary", "")
    if raw_summary:
        # 尽量安全转文本，避免 HTML, NaN, None 等问题
        try:
            text = str(raw_summary)
        except Exception:
            text = repr(raw_summary)

        safe_summary = escape(text, quote=False)
        st.markdown(
            f'<div class="stMarkdown">{safe_summary}</div>',
            unsafe_allow_html=True
        )
    else:
        st.markdown("<div class='stMarkdown'>(No summary available)</div>", unsafe_allow_html=True)

    # Week of / Publisher / Publish date
    week_of = None
    if isinstance(row.get("publish_date"), date):
        week_of = (row["publish_date"] - timedelta(days=row["publish_date"].weekday())).isoformat()
    meta = (
        (f"**Publisher:** {row.get('publisher','')}  " if row.get("publisher") else "") +
        (f"**Publish date:** {row.get('publish_date','')}  " if row.get("publish_date") else "")
    )
    if meta:
        st.markdown(meta)

    if row.get("url"):
        st.link_button("Open article ↗", row["url"], use_container_width=True)

    # --- AI 预选 = Category，规范化并勾选 ---
    st.markdown("#### Recommended Categories (AI Pre-selection)")

    def normalize_cat(s: str) -> str:
        s = (s or "").strip()
        canon = {
            "housing affordability": "Housing Affordability",
            "culture led development": "Culture Led Development",
            "culture-led development": "Culture Led Development",
            "net zero cities": "Net Zero Cities",
            "public/private development": "Public/Private Development",
            "public / private development": "Public/Private Development",
        }
        return canon.get(s.lower(), s)

    preselected = []
    if row.get("category"):
        preselected = [
            normalize_cat(x)
            for x in str(row["category"]).split(";")
            if normalize_cat(x) in CATEGORIES
        ]
    sel = set(preselected)

    cols = st.columns(2)
    with cols[0]:
        c1 = st.checkbox("Housing Affordability", value=("Housing Affordability" in sel))
        c2 = st.checkbox("Net Zero Cities", value=("Net Zero Cities" in sel))
    with cols[1]:
        c3 = st.checkbox("Public/Private Development", value=("Public/Private Development" in sel))
        c4 = st.checkbox("Culture Led Development", value=("Culture Led Development" in sel))

    selected_categories = [
        name for name, flag in [
            ("Housing Affordability", c1),
            ("Net Zero Cities", c2),
            ("Public/Private Development", c3),
            ("Culture Led Development", c4),
        ] if flag
    ]
    categories_str = "; ".join(selected_categories)

    st.divider()

    # --- 人工审核，仅写入 news_reviews ---
    st.markdown("#### Manual Review")
    decision = st.radio("Decision", ["Confirm", "Reject"], horizontal=True, index=0)
    note = st.text_area("Notes (optional)", value="", placeholder="Add reviewer notes...")

    col_save, col_sp = st.columns([1, 3])
    # 保存审核（Save Review）这段：仅替换 reviewed_at 这一行
    with col_save:
        if st.button("💾 Save Review", use_container_width=True):
            try:
                supabase.table("news_reviews").insert({
                    "title": row.get("title", ""),
                    "publisher": row.get("publisher", ""),
                    "publish_date": str(row.get("publish_date") or ""),
                    "link": row.get("url", ""),
                    "decision": decision.lower(),          # confirm / reject
                    "categories": categories_str,          # '; ' 分隔
                    "note": note,
                    "summary": row.get("summary", ""),
                    "reviewed_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
                }).execute()
                st.success("Review saved to table `news_reviews`.")

                # ✅ 如果是 Confirm，把当前这条文章的信息存到 session_state，用于下面显示模板
                if decision.lower() == "confirm":
                    st.session_state["last_confirmed"] = {
                        "row": row,
                        "categories": categories_str,
                    }

            except Exception as e:
                st.error(f"Insert to `news_reviews` failed: {e}")



# ---------------------------
# 右列：统计与外链
# ---------------------------
with right:
    st.subheader("Weekly Articles")
    st.metric("Count in view", len(df))
    cnt = df.groupby("publisher", dropna=True).size().reset_index(name="count").sort_values("count", ascending=False)
    st.dataframe(cnt.rename(columns={"publisher":"source"}), use_container_width=True, height=360)

with st.expander("📝 Generate Weekly DOCX Report", expanded=False):
    # 周一选择（默认当前周周一）
    today = date.today()
    default_monday = today - timedelta(days=today.weekday())
    monday = st.date_input("Week (pick the Monday)", value=default_monday)

    author = st.text_input("Urban Lab Author", value="Your Name", key="report_author")
    outdir = st.text_input("Save to directory", value=OUTPUT_DIR,
                           help="本地保存路径；同时会提供在线下载")

    gen_col1, gen_col2 = st.columns([1,2])
    with gen_col1:
        gen_btn = st.button("Generate DOCX", type="primary", use_container_width=True)

    if gen_btn:
        try:
            rows = _fetch_reviews_week(monday)
            if not rows:
                st.warning("本周暂无审核记录。")
            else:
                docx_bytes = build_weekly_docx(rows, monday, author)
                # 1) 在线下载
                fname = f"UrbanLab_Weekly_{monday.isoformat()}.docx"
                st.download_button("Download DOCX", data=docx_bytes, file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                # 2) 本机保存
                try:
                    os.makedirs(outdir, exist_ok=True)
                    save_path = os.path.join(outdir, fname)
                    with open(save_path, "wb") as f:
                        f.write(docx_bytes.getvalue())
                    st.success(f"已保存到：{save_path}")
                except Exception as e:
                    st.warning(f"保存到本地失败：{e}")
        except Exception as e:
            st.error(f"生成失败：{e}")

# ----------------------------------------
# 单一文章的 Summary 模板展示（格式对齐截图）
# ----------------------------------------
st.markdown("---")
st.subheader("📄 Format to use for Summary (copy & paste)")

data = st.session_state.get("last_confirmed")

if not data:
    st.info("Once you complete the Confirm action with Save above, a copy-ready Weekly Report summary will appear here.")
else:
    r = data["row"]
    cats = data["categories"]
    author_for_block = st.session_state.get("report_author")

    # 处理日期
    pubdate = r.get("publish_date")
    pub_str = ""
    try:
        if isinstance(pubdate, date):
            pub_str = pubdate.strftime("%m.%d.%Y")
        elif pubdate:
            pub_str = pd.to_datetime(pubdate).strftime("%m.%d.%Y")
    except Exception:
        pub_str = str(pubdate or "")

    link = r.get("url") or r.get("link") or ""

    from html import escape as _esc

    # 顶部文字 + 字段（Title / Source / Date / Link / Author / Article Photograph）
    top_html = f"""

    <p><b>Title:</b> {_esc(r.get('title', '') or '')}</p>
    <p><b>Source:</b> {_esc(r.get('publisher', '') or '')}</p>
    <p><b>Date Published:</b> {_esc(pub_str)}</p>
    <p><b>Link:</b> <a href="{_esc(link)}">{_esc(link)}</a></p>
    <p><b>Urban Lab Author:</b> {_esc(author_for_block)}</p>
    <p><b>Article Photograph:</b></p>
    """
    st.markdown(top_html, unsafe_allow_html=True)

    # 图片（优先 image_url，再回退 og:image）
    img_bytes = None
    img_url = (r.get("image_url") or "").strip()
    if img_url:
        img_bytes = fetch_remote_img(img_url)
    if (not img_bytes) and link:
        og_url = fetch_og_image_url_with_curl(link)
        if og_url:
            img_bytes = fetch_remote_img(og_url)

    if img_bytes:
        st.image(img_bytes, width=400)
    else:
        st.write("(No image available)")

    # Summary + Initiative（红色加粗）
    summary_text = r.get("summary", "") or ""
    bottom_html = f"""
    <p><b>Article Summary:</b> {_esc(summary_text)}</p>
    <p><span style="color: red; font-weight: bold;">
        Initiative: {_esc(cats or '')}
    </span></p>
    """
    st.markdown(bottom_html, unsafe_allow_html=True)

        # ---------- Copy to clipboard button (HTML, 保留格式) ----------
    # 这段 HTML 会被复制到剪贴板，Google Docs 会按富文本粘贴
    clipboard_html = f"""
    <p><b>Title:</b> {_esc(r.get('title', '') or '')}<br>
    <b>Source:</b> {_esc(r.get('publisher', '') or '')}<br>
    <b>Date Published:</b> {_esc(pub_str)}<br>
    <b>Link:</b> <a href="{_esc(link)}">{_esc(link)}</a><br>
    <b>Urban Lab Author:</b> {_esc(author_for_block)}<br>
    <b>Article Photograph:</b> [insert image here]</p>

    <p><b>Article Summary:</b> {_esc(summary_text)}</p>

    <p><b>Initiative:</b> <span style="color: red; font-weight: bold;">
        {_esc(cats or '')}
    </span></p>
    """

    # 避免在 JS 模板字符串里把 ` 和 </script> 搞坏
    js_safe_html = (
        clipboard_html
        .replace("\\", "\\\\")
        .replace("`", "\\`")
        .replace("</script>", "<\\/script>")
    )

    components.html(
        f"""
        <button onclick="copySummaryHtml()"
                style="margin-top:8px;padding:6px 12px;font-size:14px;">
            Copy summary
        </button>
        <script>
        async function copySummaryHtml() {{
            const html = `{js_safe_html}`;
            const type = "text/html";
            const blob = new Blob([html], {{ type }});
            const data = [new ClipboardItem({{ [type]: blob }})];
            try {{
                await navigator.clipboard.write(data);
                alert("Summary copied to clipboard with formatting.");
            }} catch (e) {{
                alert("Copy failed: " + e);
            }}
        }}
        </script>
        """,
        height=60,
    )
